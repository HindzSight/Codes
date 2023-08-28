from docx import Document
import re

wordDoc = Document('E:\Download\Timetable.docx')
doc2 = Document('E:\Download\BTech5semCSE.docx')

print(":: Time Tablerrrrrr ::\n")

tables = doc2.tables[0]

time_table_data = []

for row in tables.rows:
    row_data = []
    for col_index, cell in enumerate(row.cells):
            if col_index != 5: 
                row_data.append(cell.text)
    time_table_data.append(row_data)
    # row_data[ 'index' ] will list 1st,2nd,3rd lecture 
        # row_data = [cell.text for cell in row.cells]
        # time_table_data.append(row_data)
                
# print(len(time_table_data))
d = len(time_table_data)
print('Input data :')
for data in time_table_data:
    print(data)

# match = search(r'\b(BY|B[5-8])\b',time_table_data[1][1])
# print('\nMATCHES ::::: \n',match.group())
# print('\n')

batchBY = []

print('\nTime Schedule : \n')

patt_5 = r'\b^BY.{10}|^B5.{18}|^B1B2B3B4B5B6.{10}|9.00-9.55|10.00-10.55|11.00-11.55|12.00-12.55|2.00-2.55|3.00-3.55|4.00-4.55|Mon|Tue|Wed|Thu|Fri|Sat\b'

for row_data in time_table_data:
    found_batch = False
    batchBYrow = []
    for cell_data in row_data:
        cell_data = cell_data.replace(' ','-')
        cell_data = cell_data.replace('\n', ' ')
        # if re.search(r'\b(BY|B[5-8])\b', cell_data):
        match = re.findall(patt_5,cell_data)
        if match:
            batchBYrow.append(match)
            found_batch = True
        else :
            batchBYrow.append('-')
    batchBY.append(batchBYrow)

# Printing Time Table like a Time Table 
print('Output table after processing :')
for data in batchBY:
    print(data)