from docx import Document
import re

wordDoc = Document('Z:\Download\Timetable.docx')
doc2 = Document('Z:\Download\BTech5semCSE.docx')

print(":: Time Tablerrrrrr ::\n")

tables = doc2.tables[0]

time_table_data = []

for row in tables.rows:
    row_data = []
    for col_index, cell in enumerate(row.cells):
            if col_index != 5: 
                row_data.append(cell.text)
    time_table_data.append(row_data)
    # row_data[ 'index' ] will list 1st,2nd,3rd lecture 
        # row_data = [cell.text for cell in row.cells]
        # time_table_data.append(row_data)
                
# print(len(time_table_data))
d = len(time_table_data)
print('Input data :')
for data in time_table_data:
    print(data)

# match = search(r'\b(BY|B[5-8])\b',time_table_data[1][1])
# print('\nMATCHES ::::: \n',match.group())
# print('\n')

batchBY = []

print('\nTime Schedule : \n')
# FKKKKKKKKKKKKKKK
patt_1 = r'\b(BY|B5|9.00-9.55|10.00-10.55|11.00-11.55|12.00-12.55|2.00-2.55|3.00-3.55|4.00-4.55|Mon|Tue|Wed|Thu|Fri|Sat)\b'
patt_2 = r'\bB[Y5-8]+(-[A-Z]+)+\b'
patt_3 = r'\bBY-[A-Z]+-[A-Z]+-\w+\b'
patt_4 = r'\b((BY|B5|9.00-9.55|10.00-10.55|11.00-11.55|12.00-12.55|2.00-2.55|3.00-3.55|4.00-4.55|Mon|Tue|Wed|Thu|Fri|Sat) | (B[Y0-12]+(-[A-Z0-9]+)+))\b'
patt_5 = r'^[BYB5]\w+'
patt_L = r'\bBY.\b'
patt_ex= r'\bB5.\b'
patt_p = r'\bB1B2B3B4B5B6.\b'
patt_study = r'^((B5|B1B2B3B4B5B6|BY)|((9\.00-9\.55|10\.00-10\.55|11\.00-11\.55|12\.00-12\.55|2\.00-2\.55|3\.00-3\.55|4\.00-4\.55)|Mon|Tue|Wed|Thu|Fri|Sat))'
trial = r'((BY)|(B5)|(B1B2B3B4B5B6)|(((9\.00-9\.55|10\.00-10\.55|11\.00-11\.55|12\.00-12\.55|2\.00-2\.55|3\.00-3\.55|4\.00-4\.55)|Mon|Tue|Wed|Thu|Fri|Sat)))'
trial2 = r'(BY)|(B5)|(B1B2B3B4B5B6)|(((9\.00-9\.55|10\.00-10\.55|11\.00-11\.55|12\.00-12\.55|2\.00-2\.55|3\.00-3\.55|4\.00-4\.55)|Mon|Tue|Wed|Thu|Fri|Sat))'
trial3 = r'(BY\S+)|(B5\S+)|(B1B2B3B4B5B6\S+)|(((9\.00-9\.55|10\.00-10\.55|11\.00-11\.55|12\.00-12\.55|2\.00-2\.55|3\.00-3\.55|4\.00-4\.55)|Mon|Tue|Wed|Thu|Fri|Sat))'
lecture_pattern = r'\b(B5|BY|B1B2B3B4B5B6)\S*\b'


for row_data in time_table_data:
    batchBYrow = []
    for cell_data in row_data:
        found_batch = False
        cell_data = cell_data.replace(' ','-')
        cell_data = cell_data.replace('\n', ' ')
        # print('DATA POINT           ::: ',cell_data)
        if re.search(trial3, cell_data) :
            # match = re.search(patt_1,cell_data)
            found_batch = True
            # print(match.group())
            batchBYrow.append(cell_data)
            # print('DATA POINT AFTER     ::: ',cell_data)
    batchBY.append(batchBYrow)
    
    if not found_batch:
        batchBYrow.append('Free Period')
    
    print('\n-------------------------')

# Printing Time Table like a Time Table 
print('Output table after processing :')
for data in batchBY:
    print(data)
# print(batchBY)
    