import re
from docx import Document

wordDoc = Document('Z:\Download\Timetable.docx')
doc2 = Document('Z:\Download\BTech5semCSE.docx')

print(":: Time Tablerrrrrr ::\n")

tables = doc2.tables[0]

time_table_data = []

for row in tables.rows:
    row_data = []
    for col_index, cell in enumerate(row.cells):
        if col_index != 5:
            row_data.append(cell.text)
    time_table_data.append(row_data)

d = len(time_table_data)
print('Input data :')
for data in time_table_data:
    print(data)

# Corrected regex patterns
patt_study = r'\b(BY|B5|B5B6|B[1-6])\b'

# Initialize dictionaries to store day and time indices
day_index = {day: index for index, day in enumerate(time_table_data[1][1:])}
time_index = {time: index for index, time in enumerate(row[0] for row in time_table_data[1:])}

# Initialize the schedule grid
schedule_grid = [['' for _ in range(len(day_index))] for _ in range(len(time_index))]

for row_index, row_data in enumerate(time_table_data[1:]):
    for col_index, cell_data in enumerate(row_data[1:]):
        cell_data = cell_data.replace(' ', '-')
        cell_data = cell_data.replace('\n', ' ')

        # Use the get method with a default value for the '9.00-9.55' time value
        time = time_table_data[0][col_index + 1]
        day = row_data[0]
        schedule_grid[time_index.get(time, 0)][day_index[day]] = cell_data

# Print the schedule grid
print('\nTime Schedule : \n')
for time, row in zip(time_index.keys(), schedule_grid):
    print(f"{time} | {' | '.join(row)}")
