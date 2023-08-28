from docx import Document
import re

wordDoc = Document('E:\Download\Timetable.docx')
doc2 = Document('E:\Download\BTech5semCSE.docx')

print(":: Time Tablerrrrrr ::\n")

tables = doc2.tables[0]

time_table_data = []

for row in tables.rows:
    row_data = []
    for col_index, cell in enumerate(row.cells):
            if col_index != 5: 
                row_data.append(cell.text)
    time_table_data.append(row_data)
    # row_data[ 'index' ] will list 1st,2nd,3rd lecture 
        # row_data = [cell.text for cell in row.cells]
        # time_table_data.append(row_data)
                
# print(len(time_table_data))
d = len(time_table_data)
print('Input data :')
for data in time_table_data:
    print(data)

# match = search(r'\b(BY|B[5-8])\b',time_table_data[1][1])
# print('\nMATCHES ::::: \n',match.group())
# print('\n')

batchBY = []

print('\nTime Schedule : \n')

patt_1 = r'\b(BY|B5|9.00-9.55|10.00-10.55|11.00-11.55|12.00-12.55|2.00-2.55|3.00-3.55|4.00-4.55|Mon|Tue|Wed|Thu|Fri|Sat)\b'
patt_2 = r'\bB[Y5-8]+(-[A-Z]+)+\b'
patt_3 = r'\bBY-[A-Z]+-[A-Z]+-\w+\b'
patt_4 = r'\b((BY|B5|9.00-9.55|10.00-10.55|11.00-11.55|12.00-12.55|2.00-2.55|3.00-3.55|4.00-4.55|Mon|Tue|Wed|Thu|Fri|Sat) | (B[Y0-12]+(-[A-Z0-9]+)+))\b'
patt_5 = r'^[BYB5]\w+'
patt_L = r'\bBY.\b'
patt_ex= r'\bB5.\b'
patt_p = r'\bB1B2B3B4B5B6.\b'



for row_data in time_table_data:
    found_batch = False
    batchBYrow = []
    for cell_data in row_data:
        # if re.search(r'\b(BY|B[5-8])\b', cell_data):
        print('CELL DATA            ::: ',cell_data)
        cell_data = cell_data.replace(' ','-')
        cell_data = cell_data.replace('\n', ' ')
        print('DATA POINT           ::: \"',cell_data,'\"')
        print('SIZE : ',len(cell_data))
        if len(cell_data) < 2 :
            continue
        elif cell_data[1]  == 'Y':
            found_batch = True
            # print(cell_data)
            batchBYrow.append(cell_data)
            print('DATA POINT AFTER     ::: ',cell_data)
    batchBY.append(batchBYrow)
    
    if not found_batch:
        batchBYrow.append('Free Period')
    
    print('\n-------------------------')

# Printing Time Table like a Time Table 
print('Output table after processing :')
for data in batchBY:
    print(data)
# print(batchBY)
    
# Split the cell_data on the basis of spaces and use them to solve string problem check shell for the o/p 31/07/2023 / 6:45 PM

# Output table after processing :
# ['Free Period']
# ['BY-L-CoE-LT2 B3-T--PTRP–CR24 ', 'BY-L-AP-Lab-2-LT2 B3-T-ToC-CR17 B9-T--PTRP–CR24']
# ['Free Period']
# ['BY-L-FSD-LT2 B9-T-CoE-CR27 B10-T--ToC-CR17', 'BY-L-ToC-LT2 B5-T--PTRP–CR24']
# ['Free Period']
# ['Free Period']
# ['Free Period']
# ['BY-L-FSD-LT1 B11-T-CoE-CR27 B12-T--ToC-CR17', 'BY-L-ToC-LT2  ', 'BY-L-PTRP-LT1 B2-T-CoE-CR27 B4-T--PTRP–CR24 ']
# ['Free Period']
# ['BY-L-CoE-LT2 B2-T--PTRP–CR24 ', 'BY-L-PTRP-LT2 BZ-L-FSD-LT1 B3-T-CoE-CR27 B4-T--ToC-CR17']
# ['BY-L-PTRP-LT2 BZ-L-FSD-LT1 B3-T-CoE-CR27 B4-T--ToC-CR17']
# ['BY-L-ToC-LT2 ']
# ['Free Period']